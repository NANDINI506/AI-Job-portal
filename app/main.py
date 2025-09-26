from fastapi import FastAPI, Form, Request, HTTPException, Depends, UploadFile, Query
from fastapi.responses import HTMLResponse, Response, RedirectResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
import io
import pandas as pd
import os
import re
import json
from datetime import datetime, timedelta
from typing import Optional
from sqlalchemy.orm import Session
import hashlib
import shutil
from dotenv import load_dotenv
load_dotenv()

from .database import Base, engine, SessionLocal
from . import models, analysis, scraper
from .nlp_utils import extract_skills, extract_entities, calculate_skill_match, get_skill_suggestions
from .email_utils import send_email
from .resume_builder.resume_builder import ResumeBuilder
from .resume_builder.models import GeneratedResume
from .resume_builder.pdf_generator import ResumePDFGenerator
from apscheduler.schedulers.background import BackgroundScheduler
from apscheduler.triggers.cron import CronTrigger
from sentence_transformers import SentenceTransformer

# Create all tables
Base.metadata.create_all(bind=engine)

app = FastAPI()

# ------------------------
# Initialize sentence-transformers model
# ------------------------
MODEL_PATH = "./models/all-MiniLM-L6-v2"

if not os.path.exists(MODEL_PATH):
    model = SentenceTransformer('all-MiniLM-L6-v2')
    model.save(MODEL_PATH)
else:
    model = SentenceTransformer(MODEL_PATH)


# Get paths from environment variables with fallbacks
STATIC_DIR = os.environ.get('STATIC_DIR', 'app/static')
RESUMES_DIR = os.environ.get('RESUMES_DIR', 'resumes')
TEMPLATES_DIR = os.environ.get('TEMPLATES_DIR', 'app/templates')

# Ensure directories exist
os.makedirs(STATIC_DIR, exist_ok=True)
os.makedirs(RESUMES_DIR, exist_ok=True)

app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")
app.mount("/resumes", StaticFiles(directory=RESUMES_DIR), name="resumes")
templates = Jinja2Templates(directory=TEMPLATES_DIR)

# User session management
current_user = None
is_supervisor = False

# Database dependency
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

def hash_password(password: str) -> str:
    return hashlib.sha256(password.encode()).hexdigest()

@app.get("/register", response_class=HTMLResponse)
async def register_page(request: Request):
    print("Accessing registration page")  # Debug log
    if current_user:
        print(f"User {current_user} is logged in, redirecting to dashboard")  # Debug log
        return RedirectResponse(url="/dashboard", status_code=303)
    print("Rendering registration template")  # Debug log
    return templates.TemplateResponse("register.html", {"request": request, "current_user": current_user, "is_supervisor": is_supervisor})

@app.post("/register")
async def register(
    request: Request,
    username: str = Form(...),
    email: str = Form(...),
    password: str = Form(...),
    confirm_password: str = Form(...),
    db: Session = Depends(get_db)
):
    print(f"Registration attempt for username: {username}")  # Debug log
    try:
        if password != confirm_password:
            print("Passwords do not match")  # Debug log
            return templates.TemplateResponse("register.html", {
                "request": request,
                "error": "Passwords do not match",
                "current_user": current_user,
                "is_supervisor": is_supervisor
            })
        
        # Check if username or email already exists
        existing_user = db.query(models.User).filter(
            (models.User.username == username) | (models.User.email == email)
        ).first()
        
        if existing_user:
            print(f"User {username} or email {email} already exists")  # Debug log
            return templates.TemplateResponse("register.html", {
                "request": request,
                "error": "Username or email already exists",
                "current_user": current_user,
                "is_supervisor": is_supervisor
            })
        
        # Create new user
        hashed_password = hash_password(password)
        new_user = models.User(
            username=username,
            email=email,
            password=hashed_password,
            is_supervisor=False
        )
        
        db.add(new_user)
        db.commit()
        print(f"Successfully registered user: {username}")  # Debug log
        
        return RedirectResponse(url="/login", status_code=303)
    except Exception as e:
        print(f"Registration error: {str(e)}")  # Debug log
        db.rollback()
        return templates.TemplateResponse("register.html", {
            "request": request,
            "error": f"Registration failed: {str(e)}",
            "current_user": current_user,
            "is_supervisor": is_supervisor
        })

@app.post("/login")
async def login(
    request: Request,
    username: str = Form(...),
    password: str = Form(...),
    db: Session = Depends(get_db)
):
    global current_user, is_supervisor
    
    # Check for admin login
    if username == "admin" and password == "admin":
        current_user = username
        is_supervisor = True
        return RedirectResponse(url="/dashboard", status_code=303)
    
    # Check for regular user login
    hashed_password = hash_password(password)
    user = db.query(models.User).filter(
        models.User.username == username,
        models.User.password == hashed_password
    ).first()
    
    if user:
        current_user = user.username
        is_supervisor = user.is_supervisor
        return RedirectResponse(url="/dashboard", status_code=303)
    else:
        return templates.TemplateResponse("login.html", {
            "request": request,
            "error": "Invalid credentials",
            "current_user": current_user,
            "is_supervisor": is_supervisor
        })

@app.get("/login", response_class=HTMLResponse)
async def login_page(request: Request):
    return templates.TemplateResponse("login.html", {"request": request, "current_user": current_user, "is_supervisor": is_supervisor})

@app.get("/logout")
async def logout():
    global current_user, is_supervisor
    current_user = None
    is_supervisor = False
    return RedirectResponse(url="/login", status_code=303)

@app.get("/", response_class=HTMLResponse)
async def home(request: Request, db: Session = Depends(get_db)):
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    jobs = db.query(models.Job).all()
    return templates.TemplateResponse("dashboard.html", 
        {"request": request, "jobs": jobs, "is_supervisor": is_supervisor, "current_user": current_user})

@app.get("/search-jobs", response_class=HTMLResponse)
async def search_jobs(
    request: Request,
    keyword: Optional[str] = None,
    location: Optional[str] = None,
    experience: Optional[str] = None,
    job_type: Optional[str] = None,
    db: Session = Depends(get_db)
):
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
        
    query = db.query(models.Job)
    
    if keyword:
        keyword = keyword.lower()
        query = query.filter(
            (models.Job.title.ilike(f"%{keyword}%")) |
            (models.Job.company.ilike(f"%{keyword}%")) |
            (models.Job.skills.ilike(f"%{keyword}%")) |
            (models.Job.description.ilike(f"%{keyword}%"))
        )
    
    if location:
        location = location.lower()
        query = query.filter(models.Job.location.ilike(f"%{location}%"))
    
    if experience:
        experience = experience.lower()
        if experience == "junior":
            query = query.filter(models.Job.experience_level.ilike("%Entry-level%"))
        elif experience == "mid":
            query = query.filter(
                (models.Job.experience_level.ilike("%Mid-level%")) |
                (models.Job.experience_level.ilike("%Mid-Senior%"))
            )
        elif experience == "senior":
            query = query.filter(
                (models.Job.experience_level.ilike("%Senior%")) |
                (models.Job.experience_level.ilike("%Lead%")) |
                (models.Job.experience_level.ilike("%Manager%"))
            )
    
    if job_type:
        job_type = job_type.lower()
        query = query.filter(models.Job.job_type.ilike(f"%{job_type}%"))
    
    jobs = query.all()

    # Attach preview/full flags for templates to enable Read More toggle
    try:
        from .job_description_processor import create_structured_description
        for job in jobs:
            try:
                structured = create_structured_description(job.description or "")
                setattr(job, 'description_preview', structured.get('preview', ''))
                setattr(job, 'has_more', structured.get('has_more', False))
            except Exception:
                setattr(job, 'description_preview', (job.description or '')[:300])
                setattr(job, 'has_more', bool(job.description and len(job.description) > 300))
    except Exception:
        for job in jobs:
            setattr(job, 'description_preview', (job.description or '')[:300])
            setattr(job, 'has_more', bool(job.description and len(job.description) > 300))
    
    return templates.TemplateResponse("search_jobs.html", {
        "request": request,
        "jobs": jobs,
        "searched": bool(keyword or location or experience or job_type),
        "is_supervisor": is_supervisor,
        "current_user": current_user
    })

@app.get("/add-job", response_class=HTMLResponse)
async def add_job_form(request: Request):
    if not current_user or not is_supervisor:
        raise HTTPException(status_code=403, detail="Only supervisors can add jobs")
    return templates.TemplateResponse("add_job.html", {"request": request, "current_user": current_user, "is_supervisor": is_supervisor})

@app.post("/add-job")
async def add_job(
    request: Request,
    title: str = Form(...),
    company: str = Form(...),
    location: str = Form(...),
    description: str = Form(""),
    salary: float = Form(...),
    skills: str = Form(...),
    experience_level: str = Form(...),
    job_type: str = Form(...),
    remote_policy: str = Form(...),
    db: Session = Depends(get_db)
):
    if not current_user or not is_supervisor:
        raise HTTPException(status_code=403, detail="Only supervisors can add jobs")
    # Validation: Require non-empty, non-trivial description and skills
    if not description or len(description.strip()) < 10 or not skills or len(skills.strip()) < 3:
        return templates.TemplateResponse("add_job.html", {
            "request": request,
            "current_user": current_user,
            "is_supervisor": is_supervisor,
            "error": "Please provide a full job description (at least 10 characters) and at least one skill."
        })
    new_job = models.Job(
        title=title,
        company=company,
        location=location,
        description=description,
        salary=salary,
        skills=skills,
        experience_level=experience_level,
        job_type=job_type,
        remote_policy=remote_policy,
        posted_by=current_user
    )
    
    db.add(new_job)
    db.commit()
    db.refresh(new_job)

    # Email notification logic
    try:
        from .nlp_utils import extract_skills, calculate_skill_match
        job_text = (description or "") + " " + (skills or "")
        job_skills = extract_skills(job_text)
        users = db.query(models.User).filter(models.User.is_supervisor == False).all()
        for user in users:
            user_skills_str = user.skills or ""
            if user_skills_str.strip():
                user_skills_text = user_skills_str.replace(",", " ")
                user_skills = extract_skills(user_skills_text)
                match_scores = calculate_skill_match(job_skills, user_skills)
                score = match_scores.get('overall', 0)
                # Lower threshold to 30% to catch more matches
                if score > 30 and user.email:
                    subject = f"New Job Match: {title} at {company}"
                    desc_snippet = (description[:150] + '...') if description and len(description) > 150 else (description or '')
                    job_url = f"{request.base_url}analyze_job_match/{new_job.id}"
                    body = f"""
                    <p>Dear {user.full_name or user.username},</p>
                    <p>A new job that matches your skills has been posted:</p>
                    <ul>
                        <li><b>Title:</b> {title}</li>
                        <li><b>Company:</b> {company}</li>
                        <li><b>Location:</b> {location}</li>
                        <li><b>Description:</b> {desc_snippet}</li>
                        <li><b>Required Skills:</b> {skills}</li>
                    </ul>
                    <p>Your match score: <b>{score:.1f}%</b></p>
                    <p><a href=\"{job_url}\">View this job and apply now!</a></p>
                    <p>Best regards,<br>Job Portal Team</p>
                    """
                    send_email(user.email, subject, body)
    except Exception as e:
        print(f"[Email Notification Error] {e}")

    return RedirectResponse(url="/dashboard", status_code=303)

@app.post("/apply-job/{job_id}")
async def apply_job(
    request: Request,
    job_id: int,
    cover_letter: str = Form(...),
    resume_option: str = Form("upload"),
    resume: Optional[UploadFile] = Form(None),
    ai_resume_id: Optional[int] = Form(None),
    rating: Optional[int] = Form(None),
    comment: Optional[str] = Form(None),
    db: Session = Depends(get_db)
):
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    # Get user from database
    user = db.query(models.User).filter(models.User.username == current_user).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    # Check if job exists
    job = db.query(models.Job).filter(models.Job.id == job_id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    # Handle resume upload/selection
    resume_filename = None
    
    if resume_option == "upload" and resume:
        # Handle new resume upload
        if resume.filename:
            # Save the uploaded file
            file_extension = resume.filename.split('.')[-1].lower()
            if file_extension not in ['pdf', 'doc', 'docx']:
                raise HTTPException(status_code=400, detail="Invalid file type. Only PDF, DOC, and DOCX files are allowed.")
            
            # Generate unique filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            resume_filename = f"resume_{user.id}_{timestamp}.{file_extension}"
            file_path = os.path.join(RESUMES_DIR, resume_filename)
            
            # Save file
            with open(file_path, "wb") as buffer:
                shutil.copyfileobj(resume.file, buffer)
            
            # Update user's resume filename
            user.resume_filename = resume_filename
            db.commit()
    
    elif resume_option == "existing":
        # Use existing resume
        resume_filename = user.resume_filename
    
    elif resume_option == "ai" and ai_resume_id:
        # Use AI generated resume
        ai_resume = db.query(GeneratedResume).filter(
            GeneratedResume.id == ai_resume_id,
            GeneratedResume.user_id == user.id
        ).first()
        if ai_resume:
            resume_filename = f"ai_resume_{ai_resume_id}.pdf"
    
    # Create new application
    new_application = models.Application(
        job_id=job_id,
        user_id=user.id,
        cover_letter=cover_letter
    )
    db.add(new_application)
    db.commit()
    db.refresh(new_application)
    # Create review if provided
    if rating and comment and str(rating).isdigit():
        new_review = models.Review(
            user_id=user.id,
            job_id=job_id,
            rating=int(rating),
            comment=comment
        )
        db.add(new_review)
        db.commit()
    return RedirectResponse(url=f"/analyze_job_match/{job_id}", status_code=303)

@app.get("/apply_job/{job_id}", response_class=HTMLResponse)
async def apply_job_page(
    request: Request,
    job_id: int,
    db: Session = Depends(get_db)
):
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    
    # Get user from database with generated resumes
    user = db.query(models.User).filter(models.User.username == current_user).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    # Fetch user's generated resumes
    from .resume_builder.models import GeneratedResume
    user.generated_resumes = db.query(GeneratedResume).filter(GeneratedResume.user_id == user.id).all()
    print(f"Found {len(user.generated_resumes)} generated resumes for user {user.username}")
    
    # Check if job exists
    job = db.query(models.Job).filter(models.Job.id == job_id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    
    # Check if user has already applied
    existing_application = db.query(models.Application).filter(
        models.Application.job_id == job_id,
        models.Application.user_id == user.id
    ).first()
    
    # Prepare preview for apply page description (Read More toggle)
    try:
        from .job_description_processor import create_structured_description
        structured = create_structured_description(job.description or "")
        setattr(job, 'description_preview', structured.get('preview', ''))
        setattr(job, 'has_more', structured.get('has_more', False))
    except Exception:
        setattr(job, 'description_preview', (job.description or '')[:300])
        setattr(job, 'has_more', bool(job.description and len(job.description) > 300))
    
    return templates.TemplateResponse("apply_job.html", {
        "request": request,
        "job": job,
        "user": user,
        "applied": existing_application is not None,
        "current_user": current_user,
        "is_supervisor": is_supervisor
    })

@app.post("/delete-job/{job_id}")
async def delete_job(
    request: Request, 
    job_id: int,
    db: Session = Depends(get_db)
):
    print(f"Delete request received for job_id: {job_id}")
    print(f"Current user: {current_user}")
    print(f"Is supervisor: {is_supervisor}")
    
    if not current_user or not is_supervisor:
        print("Access denied: Not a supervisor")
        raise HTTPException(status_code=403, detail="Only supervisors can delete jobs")
        
    job = db.query(models.Job).filter(models.Job.id == job_id).first()
    if not job:
        print(f"Job not found with id: {job_id}")
        raise HTTPException(status_code=404, detail="Job not found")
    
    print(f"Deleting job: {job.title} (ID: {job.id})")
    db.delete(job)
    db.commit()
    print("Job deleted successfully")
    
    return RedirectResponse(url="/dashboard", status_code=303)

@app.get("/edit-job/{job_id}", response_class=HTMLResponse)
async def edit_job_form(
    request: Request, 
    job_id: int,
    db: Session = Depends(get_db)
):
    if not current_user or not is_supervisor:
        raise HTTPException(status_code=403, detail="Only supervisors can edit jobs")
        
    job = db.query(models.Job).filter(models.Job.id == job_id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
        
    return templates.TemplateResponse("edit_job.html", {
        "request": request,
        "job": job,
        "job_id": job_id,
        "is_supervisor": is_supervisor,
        "current_user": current_user
    })

@app.post("/edit-job/{job_id}")
async def edit_job(
    request: Request,
    job_id: int,
    title: str = Form(...),
    company: str = Form(...),
    location: str = Form(...),
    description: str = Form(""),
    salary: float = Form(...),
    skills: str = Form(...),
    experience_level: str = Form(...),
    job_type: str = Form(...),
    remote_policy: str = Form(...),
    db: Session = Depends(get_db)
):
    if not current_user or not is_supervisor:
        raise HTTPException(status_code=403, detail="Only supervisors can edit jobs")
        
    job = db.query(models.Job).filter(models.Job.id == job_id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
        
    job.title = title
    job.company = company
    job.location = location
    job.description = description
    job.salary = salary
    job.skills = skills
    job.experience_level = experience_level
    job.job_type = job_type
    job.remote_policy = remote_policy
    
    db.commit()
    db.refresh(job)
    
    return RedirectResponse(url="/dashboard", status_code=303)

def get_jobs(db: Session = Depends(get_db)):
    return db.query(models.Job).all()

@app.get("/plot/skills")
async def get_skills_plot(db: Session = Depends(get_db)):
    jobs = db.query(models.Job).all()
    # Convert jobs to DataFrame for analysis
    df = pd.DataFrame([{
        'title': job.title,
        'company': job.company,
        'location': job.location,
        'salary': job.salary,
        'skills': job.skills
    } for job in jobs])
    
    fig = analysis.create_skills_plot(df)
    return fig_to_response(fig)

@app.get("/plot/salary")
async def get_salary_plot(db: Session = Depends(get_db)):
    jobs = db.query(models.Job).all()
    # Convert jobs to DataFrame for analysis
    df = pd.DataFrame([{
        'title': job.title,
        'company': job.company,
        'location': job.location,
        'salary': job.salary,
        'skills': job.skills
    } for job in jobs])
    
    fig = analysis.create_salary_plot(df)
    return fig_to_response(fig)

@app.get("/dashboard", response_class=HTMLResponse)
async def dashboard(request: Request, db: Session = Depends(get_db), rec_page: int = Query(1, alias="rec_page"), admin_page: int = Query(1, alias="admin_page")):
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    jobs = get_jobs(db)
    user = db.query(models.User).filter(models.User.username == current_user).first()

    # Recommend jobs for regular users (not supervisors)
    recommended_jobs = []
    if user and not is_supervisor:
        from .nlp_utils import extract_skills, calculate_skill_match
        # Only recommend if user has uploaded a resume and has skills
        user_skills_str = user.skills or ""
        if user_skills_str.strip():
            user_skills_text = user_skills_str.replace(",", " ")
            user_skills = extract_skills(user_skills_text)
            job_scores = []
            for job in jobs:
                job_text = (job.description or "") + " " + (job.skills or "")
                job_skills = extract_skills(job_text)
                match_scores = calculate_skill_match(job_skills, user_skills)
                score = match_scores.get('overall', 0)
                if score > 50:
                    job_scores.append((job, score))
            # Sort by newest first, then by score descending
            job_scores.sort(key=lambda x: (x[0].posted_date or datetime.min, x[1]), reverse=True)
            # Pagination logic
            per_page = 6
            total = len(job_scores)
            total_pages = (total + per_page - 1) // per_page
            rec_page_clamped = max(1, min(rec_page, total_pages)) if total_pages > 0 else 1
            start = (rec_page_clamped - 1) * per_page
            end = start + per_page
            paginated_recommended_jobs = job_scores[start:end]
        else:
            paginated_recommended_jobs = []
            total_pages = 1
            rec_page_clamped = 1
    else:
        paginated_recommended_jobs = []
        total_pages = 1
        rec_page_clamped = 1

    # Admin jobs - show all on one page
    all_jobs = []
    all_jobs_total_pages = 1
    admin_page_clamped = 1
    if is_supervisor:
        all_jobs = jobs  # Show all jobs without pagination

    return templates.TemplateResponse("dashboard.html", {
        "request": request,
        "jobs": jobs,
        "is_supervisor": is_supervisor,
        "current_user": current_user,
        "user": user,
        "recommended_jobs": paginated_recommended_jobs,
        "rec_page": rec_page_clamped,
        "rec_total_pages": total_pages,
        "all_jobs": all_jobs,
        "admin_page": admin_page_clamped,
        "all_jobs_total_pages": all_jobs_total_pages
    })

@app.get("/api/jobs/latest")
async def api_latest_job(db: Session = Depends(get_db)):
    """Return latest posted_date and total count to detect updates on the client."""
    if not current_user:
        raise HTTPException(status_code=401, detail="Not authenticated")
    latest = db.query(models.Job).order_by(models.Job.posted_date.desc()).first()
    total = db.query(models.Job).count()
    return {
        "latest_posted_date": latest.posted_date.isoformat() if latest else None,
        "total": total,
    }

@app.get("/api/recommended-jobs")
async def api_recommended_jobs(
    request: Request,
    page: int = Query(1),
    per_page: int = Query(6),
    db: Session = Depends(get_db)
):
    """Return the current user's recommended jobs with match scores (top 6)."""
    if not current_user:
        raise HTTPException(status_code=401, detail="Not authenticated")
    if is_supervisor:
        return {"jobs": []}
    user = db.query(models.User).filter(models.User.username == current_user).first()
    if not user or not (user.skills or "").strip():
        return {"jobs": []}

    jobs = db.query(models.Job).all()
    from .nlp_utils import extract_skills, calculate_skill_match
    user_skills = extract_skills(user.skills.replace(",", " "))
    scored = []
    for job in jobs:
        job_text = (job.description or "") + " " + (job.skills or "")
        job_skills = extract_skills(job_text)
        score = calculate_skill_match(job_skills, user_skills).get("overall", 0)
        if score > 50:
            scored.append((job, score))
    # Newest first, then score
    scored.sort(key=lambda x: ((x[0].posted_date or datetime.min), x[1]), reverse=True)
    total = len(scored)
    total_pages = (total + per_page - 1) // per_page if per_page > 0 else 1
    page_clamped = max(1, min(page, total_pages)) if total_pages > 0 else 1
    start = (page_clamped - 1) * per_page
    end = start + per_page
    top = scored[start:end]
    return {
        "jobs": [
            {
                "id": j.id,
                "title": j.title,
                "company": j.company,
                "location": j.location,
                "salary": j.salary,
                "skills": j.skills or "",
                "posted_date": j.posted_date.isoformat() if j.posted_date else None,
                "score": float(f"{s:.2f}")
            }
            for j, s in top
        ],
        "page": page_clamped,
        "per_page": per_page,
        "total": total,
        "total_pages": total_pages
    }

@app.get("/view-applications", response_class=HTMLResponse)
async def view_applications(
    request: Request,
    db: Session = Depends(get_db)
):
    if not current_user or not is_supervisor:
        raise HTTPException(status_code=403, detail="Only supervisors can view applications")
    # Get all jobs and their applications
    jobs = db.query(models.Job).all()
    applications = db.query(models.Application).all()
    # Fetch user details for each application
    user_map = {}
    for app in applications:
        user = db.query(models.User).filter(models.User.id == app.user_id).first()
        if user:
            # Robust fallback for name and email
            name = user.full_name or user.username or user.email or 'Unknown'
            email = user.resume_email or user.email or 'Unknown'
            phone = user.resume_phone or ''
            user_map[app.id] = {
                "name": name,
                "email": email,
                "phone": phone,
                "is_admin": user.is_supervisor if hasattr(user, 'is_supervisor') else False
            }
        else:
            user_map[app.id] = {
                "name": "Unknown",
                "email": "",
                "phone": "",
                "is_admin": False
            }
    print("DEBUG user_map:", user_map)
    return templates.TemplateResponse("view_applications.html", {
        "request": request,
        "jobs": jobs,
        "applications": applications,
        "user_map": user_map,
        "is_supervisor": is_supervisor,
        "current_user": current_user
    })

@app.post("/delete-application/{application_id}")
async def delete_application(
    request: Request,
    application_id: int,
    db: Session = Depends(get_db)
):
    if not current_user or not is_supervisor:
        raise HTTPException(status_code=403, detail="Only supervisors can delete applications")
        
    application = db.query(models.Application).filter(models.Application.id == application_id).first()
    if not application:
        raise HTTPException(status_code=404, detail="Application not found")
        
    db.delete(application)
    db.commit()
    
    return RedirectResponse(url="/view-applications", status_code=303)

@app.post("/update-application-status/{application_id}")
async def update_application_status(
    request: Request,
    application_id: int,
    status: str = Form(...),
    db: Session = Depends(get_db)
):
    if not current_user or not is_supervisor:
        raise HTTPException(status_code=403, detail="Only supervisors can update application status")
        
    application = db.query(models.Application).filter(models.Application.id == application_id).first()
    if not application:
        raise HTTPException(status_code=404, detail="Application not found")
        
    application.status = status
    db.commit()
    db.refresh(application)
    
    return RedirectResponse(url="/view-applications", status_code=303)

@app.get("/analysis", response_class=HTMLResponse)
async def analysis_page(request: Request):
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    return templates.TemplateResponse("analysis.html", {
        "request": request,
        "current_user": current_user,
        "is_supervisor": is_supervisor
    })

@app.get("/api/analysis-data")
async def get_analysis_data(db: Session = Depends(get_db)):
    if not current_user:
        raise HTTPException(status_code=401, detail="Not authenticated")
    # Supervisors see all applications, users see only their own
    if is_supervisor:
        applications = db.query(models.Application).all()
    else:
        user = db.query(models.User).filter(models.User.username == current_user).first()
        applications = db.query(models.Application).filter(models.Application.user_id == user.id).all()
    dates = {}
    for app in applications:
        date = app.application_date.strftime('%Y-%m-%d')
        dates[date] = dates.get(date, 0) + 1
    # Sort dates chronologically
    sorted_dates = sorted(dates.items())
    dates_list = [date for date, _ in sorted_dates]
    counts_list = [count for _, count in sorted_dates]
    # Get most applied jobs with user information
    job_applications = {}
    for app in applications:
        job = db.query(models.Job).filter(models.Job.id == app.job_id).first()
        if job:
            key = f"{job.title}"
            job_applications[key] = job_applications.get(key, 0) + 1
    # Sort jobs by application count and get top 5
    sorted_jobs = sorted(job_applications.items(), key=lambda x: x[1], reverse=True)[:5]
    top_jobs = [job for job, _ in sorted_jobs]
    top_job_counts = [count for _, count in sorted_jobs]
    # Get job status distribution (for all jobs, or just user's if you want)
    jobs = db.query(models.Job).all() if is_supervisor else [db.query(models.Job).filter(models.Job.id == app.job_id).first() for app in applications]
    statuses = {}
    for job in jobs:
        if job:
            status = job.status if hasattr(job, 'status') else 'Active'
            statuses[status] = statuses.get(status, 0) + 1
    return {
        "dates": dates_list,
        "application_counts": counts_list,
        "top_jobs": top_jobs,
        "top_job_counts": top_job_counts,
        "statuses": list(statuses.keys()),
        "status_counts": list(statuses.values())
    }

def fig_to_response(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format='png')
    buf.seek(0)
    return Response(content=buf.getvalue(), media_type='image/png')

@app.get("/scrape-jobs", response_class=HTMLResponse)
async def scrape_jobs_page(request: Request):
    if not current_user or not is_supervisor:
        raise HTTPException(status_code=403, detail="Only supervisors can scrape jobs")
    return templates.TemplateResponse("scrape_jobs.html", {"request": request, "current_user": current_user, "is_supervisor": is_supervisor})

@app.post("/scrape-jobs")
async def scrape_jobs(
    request: Request,
    job_title: str = Form(...),
    location: str = Form(...),
    num_pages: int = Form(1),
    db: Session = Depends(get_db)
):
    if not current_user or not is_supervisor:
        raise HTTPException(status_code=403, detail="Only supervisors can scrape jobs")
    
    try:
        # Initialize progress tracking
        progress = {
            "total_jobs": 0,
            "new_jobs": 0,
            "duplicate_jobs": 0,
            "errors": 0,
            "current_page": 0
        }
        
        # Scrape jobs with progress tracking
        jobs = []
        try:
            jobs = scraper.scrape_linkedin_jobs(job_title, location, num_pages)
            progress["total_jobs"] = len(jobs)
        except Exception as e:
            return templates.TemplateResponse("scrape_jobs.html", {
                "request": request,
                "error": f"Failed to scrape jobs: {str(e)}"
            })
        
        # Add unique jobs to database with progress tracking
        new_jobs = []  # Collect all new jobs for summary notification
        for job in jobs:
            try:
                # Skip jobs with missing or too-short description or skills
                if not job.get('description') or len(job['description'].strip()) < 10 or not job.get('skills') or len(job['skills'].strip()) < 3:
            
                    continue
                # Normalize job data for better duplicate detection
                normalized_title = scraper.normalize_text(job['title'])
                normalized_company = scraper.normalize_text(job['company'])
                normalized_location = scraper.normalize_text(job['location'])
                
                # Check if job already exists using normalized data
                existing_job = db.query(models.Job).filter(
                    models.Job.title.ilike(f"%{normalized_title}%"),
                    models.Job.company.ilike(f"%{normalized_company}%"),
                    models.Job.location.ilike(f"%{normalized_location}%")
                ).first()
                
                if existing_job:
                    progress["duplicate_jobs"] += 1
                    print(f"Skipping duplicate job: {job['title']} at {job['company']}")
                    continue
                
                # Additional fuzzy matching check
                all_existing_jobs = db.query(models.Job).all()
                for existing in all_existing_jobs:
                    if scraper.is_duplicate_job(job, {
                        'title': existing.title,
                        'company': existing.company,
                        'location': existing.location
                    }):
                        progress["duplicate_jobs"] += 1
                        print(f"Skipping fuzzy duplicate: {job['title']} at {job['company']}")
                        break
                else:
                    # Create new job
                    new_job = models.Job(
                        title=job['title'],
                        company=job['company'],
                        location=job['location'],
                        description=job.get('description', ''),
                        salary=job['salary'],
                        skills=job['skills'],
                        experience_level=job.get('experience_level', 'Mid-Senior'),
                        job_type=job.get('job_type', 'Full-time'),
                        remote_policy=job.get('remote_policy', 'On-site'),
                        posted_by=current_user
                    )
                    db.add(new_job)
                    db.flush()  # Ensure new_job.id is assigned
                    db.refresh(new_job)  # Ensure new_job.id is available
                    progress["new_jobs"] += 1
                    print(f"Added new job: {job['title']} at {job['company']}")
                    new_jobs.append(new_job)  # Collect for summary notification
            except Exception as e:
                progress["errors"] += 1
                print(f"Error processing job {job.get('title', 'Unknown')}: {str(e)}")
                continue

        # Use the same email notification logic as automatic scheduler
        if new_jobs:
            try:
                from .scraper import notify_users_of_new_jobs
                notify_users_of_new_jobs(db, new_jobs)
                print(f"[Manual Scraping] Sent email notifications for {len(new_jobs)} new jobs")
            except Exception as e:
                print(f"[Manual Scraping Email Error] {e}")
        
        try:
            db.commit()
        except Exception as e:
            db.rollback()
            # Check if it's a duplicate constraint violation
            if "unique_job" in str(e).lower():
                return templates.TemplateResponse("scrape_jobs.html", {
                    "request": request,
                    "error": "Some jobs were duplicates and could not be added. Please check the logs for details."
                })
            return templates.TemplateResponse("scrape_jobs.html", {
                "request": request,
                "error": f"Failed to save jobs to database: {str(e)}"
            })
        
        # Prepare success message with detailed statistics
        message = (
            f"Scraping completed successfully!\n\n"
            f"📊 Statistics:\n"
            f"• Total jobs found: {progress['total_jobs']}\n"
            f"• New jobs added: {progress['new_jobs']}\n"
            f"• Duplicate jobs skipped: {progress['duplicate_jobs']}\n"
        )
        if progress["errors"] > 0:
            message += f"• Errors encountered: {progress['errors']}\n"
        
        message += f"\n✅ Duplicate prevention is working correctly!"
        
        return templates.TemplateResponse("scrape_jobs.html", {
            "request": request,
            "success": True,
            "message": message
        })
        
    except KeyboardInterrupt:
        db.rollback()
        return templates.TemplateResponse("scrape_jobs.html", {
            "request": request,
            "error": "Scraping was interrupted. Please try again."
        })
    except Exception as e:
        db.rollback()
        return templates.TemplateResponse("scrape_jobs.html", {
            "request": request,
            "error": f"An unexpected error occurred: {str(e)}"
        })

def auto_scrape_and_notify():
    from .scraper import scrape_linkedin_jobs, notify_users_of_new_jobs
    from .database import SessionLocal
    from . import models
    from datetime import datetime, timedelta
    
    # Check if we've already run recently (within last 2 hours) to prevent duplicate notifications
    db = SessionLocal()
    two_hours_ago = datetime.now() - timedelta(hours=2)
    
    # Check if there are any recent jobs to avoid duplicate notifications
    recent_jobs = db.query(models.Job).filter(
        models.Job.posted_date >= two_hours_ago
    ).first()
    
    if recent_jobs:
        print(f"[Auto Scraper] Skipping run - recent jobs found within 2 hours")
        db.close()
        return
    
    # Clean up old jobs (older than 30 days)
    cleanup_old_jobs(db)
    
    # More diverse job titles/locations to get better matches
    job_titles_locations = [
        ("Software Engineer", "delhi"),
        ("Full Stack Developer", "Chandigarh"),
        ("Python Developer", "delhi"),
        ("React Developer", "delhi"),
        ("Java Developer", "mohali"),
        ("Web Developer", "chandigarh"),
        ("JavaScript Developer", "delhi"),
        ("Frontend Developer", "Chandigarh"),
        ("Backend Developer", "delhi"),
        ("Node.js Developer", "mohali"),
    ]
    
    new_jobs = []
    max_jobs_per_search = 5  # Limit to 3 jobs per search (10 searches × 5 jobs = 50 max)
    
    for job_title, location in job_titles_locations:
        if len(new_jobs) >= 30:  # Stop if we already have 30 jobs
            break
            
        jobs = scrape_linkedin_jobs(job_title, location, num_pages=1)  # Reduced to 1 page
        jobs_added = 0
        
        for job in jobs:
            if jobs_added >= max_jobs_per_search:  # Stop if we have enough jobs from this search
                break
                
            if not job.get('description') or len(job['description'].strip()) < 10 or not job.get('skills') or len(job['skills'].strip()) < 3:
                continue
            existing = db.query(models.Job).filter(
                models.Job.title == job['title'],
                models.Job.company == job['company'],
                models.Job.location == job['location']
            ).first()
            if not existing:
                new_job = models.Job(
                    title=job['title'],
                    company=job['company'],
                    location=job['location'],
                    description=job['description'],
                    salary=job['salary'],
                    skills=job['skills'],
                    job_url=job.get('url', ''),  # Save the job URL
                    experience_level=job.get('experience_level', 'Mid-Senior'),
                    job_type=job.get('job_type', 'Full-time'),
                    remote_policy=job.get('remote_policy', 'On-site'),
                    posted_by="AutoScraper"
                )
                db.add(new_job)
                db.commit()
                db.refresh(new_job)
                new_jobs.append(new_job)
                jobs_added += 1
    
    if new_jobs:
        try:
            # Send email notifications
            from .scraper import notify_users_of_new_jobs
            notify_users_of_new_jobs(db, new_jobs)
            print(f"[Auto Scraping] Sent email notifications for {len(new_jobs)} new jobs")
            
            # Create in-app notifications
            create_in_app_notifications(db, new_jobs)
        except Exception as e:
            print(f"[Auto Scraping Error] {e}")
    
    db.close()

def cleanup_old_jobs(db):
    """Delete jobs older than 30 days to keep database clean"""
    from datetime import datetime, timedelta
    thirty_days_ago = datetime.now() - timedelta(days=30)
    
    # Get old jobs
    old_jobs = db.query(models.Job).filter(
        models.Job.posted_date < thirty_days_ago
    ).all()
    
    if old_jobs:
        # Delete related applications first
        for job in old_jobs:
            db.query(models.Application).filter(
                models.Application.job_id == job.id
            ).delete()
            
            # Delete related reviews
            db.query(models.Review).filter(
                models.Review.job_id == job.id
            ).delete()
            
            # Delete related notifications
            db.query(models.Notification).filter(
                models.Notification.related_job_id == job.id
            ).delete()
        
        # Delete old jobs
        deleted_count = db.query(models.Job).filter(
            models.Job.posted_date < thirty_days_ago
        ).delete()
        
        db.commit()
        print(f"Cleaned up {deleted_count} old jobs and their related data")

def create_in_app_notifications(db, new_jobs):
    """Create in-app notifications for new jobs"""
    from .nlp_utils import extract_skills, calculate_skill_match
    
    # Get all users
    users = db.query(models.User).all()
    
    for user in users:
        if not user.skills:
            continue
            
        # Calculate skill matches for this user
        user_skills_text = user.skills.replace(",", " ")
        user_skills = extract_skills(user_skills_text)
        
        matching_jobs = []
        for job in new_jobs:
            job_text = (job.description or "") + " " + (job.skills or "")
            job_skills = extract_skills(job_text)
            match_scores = calculate_skill_match(job_skills, user_skills)
            score = match_scores.get('overall', 0)
            
            if score > 30:  # Lower threshold for in-app notifications
                matching_jobs.append((job, score))
        
        if matching_jobs:
            # Sort by match score and take top 5
            matching_jobs.sort(key=lambda x: x[1], reverse=True)
            top_jobs = matching_jobs[:5]
            
            # Create notification
            job_titles = [job.title for job, _ in top_jobs]
            notification = models.Notification(
                user_id=user.id,
                title="New Jobs Match Your Skills! 🎯",
                message=f"Found {len(top_jobs)} new jobs that match your skills: {', '.join(job_titles[:3])}{' and more' if len(job_titles) > 3 else ''}",
                notification_type="job_alert",
                is_read=False
            )
            db.add(notification)
    
    db.commit()
    print(f"Created in-app notifications for {len(new_jobs)} new jobs")

# Start APScheduler when FastAPI starts
scheduler = BackgroundScheduler()
scheduler.add_job(auto_scrape_and_notify, CronTrigger(hour=11, minute=30))
scheduler.add_job(auto_scrape_and_notify, CronTrigger(hour=16, minute=0))
scheduler.start()

@app.get("/new-jobs", response_class=HTMLResponse)
async def new_jobs_page(request: Request, db: Session = Depends(get_db)):
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    
    # Get recently added jobs - look for jobs added in the last 24 hours first, then last 7 days
    from datetime import datetime, timedelta
    now = datetime.now()
    day_ago = now - timedelta(days=1)
    week_ago = now - timedelta(days=7)
    
    # First, try to get jobs from the last 24 hours (most recent scraping)
    recent_jobs = db.query(models.Job).filter(
        models.Job.posted_date >= day_ago
    ).order_by(models.Job.posted_date.desc()).all()
    
    # If no recent jobs, fall back to last 7 days
    if not recent_jobs:
        recent_jobs = db.query(models.Job).filter(
            models.Job.posted_date >= week_ago
        ).order_by(models.Job.posted_date.desc()).all()
    
    # If still no jobs, get the 20 most recently added jobs regardless of date
    if not recent_jobs:
        recent_jobs = db.query(models.Job).order_by(models.Job.posted_date.desc()).limit(20).all()
    
    # Prepare preview/full descriptions for UI toggling
    try:
        from .job_description_processor import create_structured_description
        for job in recent_jobs:
            try:
                structured = create_structured_description(job.description or "")
                # Attach ephemeral attributes for templates (not persisted)
                setattr(job, 'description_preview', structured.get('preview', ''))
                setattr(job, 'has_more', structured.get('has_more', False))
            except Exception:
                setattr(job, 'description_preview', (job.description or '')[:300])
                setattr(job, 'has_more', bool(job.description and len(job.description) > 300))
    except Exception:
        # If processor import fails, fall back gracefully
        for job in recent_jobs:
            setattr(job, 'description_preview', (job.description or '')[:300])
            setattr(job, 'has_more', bool(job.description and len(job.description) > 300))

    # Get user for skill matching
    user = db.query(models.User).filter(models.User.username == current_user).first()
    
    # Calculate match scores for each job
    job_matches = []
    if user and user.skills:
        from .nlp_utils import extract_skills, calculate_skill_match
        user_skills_text = user.skills.replace(",", " ")
        user_skills = extract_skills(user_skills_text)
        
        for job in recent_jobs:
            job_text = (job.description or "") + " " + (job.skills or "")
            job_skills = extract_skills(job_text)
            match_scores = calculate_skill_match(job_skills, user_skills)
            score = match_scores.get('overall', 0)
            job_matches.append((job, score))
    else:
        job_matches = [(job, 0) for job in recent_jobs]
    
    # Sort by posted_date (newest first) then by match score
    job_matches.sort(key=lambda x: (x[0].posted_date, x[1]), reverse=True)
    
    # Get last scraper run info
    last_scrape_info = "Today"
    if recent_jobs:
        latest_job_date = max(job.posted_date for job in recent_jobs)
        if latest_job_date.date() == now.date():
            last_scrape_info = "Today"
        elif latest_job_date.date() == (now - timedelta(days=1)).date():
            last_scrape_info = "Yesterday"
        else:
            last_scrape_info = latest_job_date.strftime("%B %d")
    
    return templates.TemplateResponse("new_jobs.html", {
        "request": request,
        "job_matches": job_matches,
        "current_user": current_user,
        "is_supervisor": is_supervisor,
        "user": user,
        "last_scrape_info": last_scrape_info,
        "now": now,
        "timedelta": timedelta
    })

@app.get("/test-scraper")
async def test_scraper():
    """Test endpoint to manually trigger the scraper"""
    try:
        # Check if we've run recently to prevent duplicate notifications
        from datetime import datetime, timedelta
        db = SessionLocal()
        two_hours_ago = datetime.now() - timedelta(hours=2)
        recent_jobs = db.query(models.Job).filter(
            models.Job.posted_date >= two_hours_ago
        ).first()
        db.close()
        
        if recent_jobs:
            return {"message": "Scraper test skipped - recent jobs found within 2 hours to prevent duplicate notifications"}
        
        auto_scrape_and_notify()
        return {"message": "Scraper test completed successfully! Check the logs for details."}
    except Exception as e:
        return {"error": f"Scraper test failed: {str(e)}"}



@app.get("/upload_resume", response_class=HTMLResponse)
async def upload_resume_page(request: Request, db: Session = Depends(get_db)):
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    user = db.query(models.User).filter(models.User.username == current_user).first()
    
    # Calculate ATS score for existing resume if available
    ats_score_data = None
    if user and user.resume_filename:
        try:
            # Read the resume file
            resume_path = os.path.join(RESUMES_DIR, user.resume_filename)
            if os.path.exists(resume_path):
                resume_text = None
                file_ext = os.path.splitext(user.resume_filename)[1].lower()
                
                if file_ext == ".pdf":
                    import pdfplumber
                    with pdfplumber.open(resume_path) as pdf:
                        resume_text = " ".join(page.extract_text() or '' for page in pdf.pages)
                elif file_ext == ".docx":
                    import docx
                    doc = docx.Document(resume_path)
                    resume_text = "\n".join([para.text for para in doc.paragraphs])
                
                if resume_text:
                    from .ats_utils import calculate_ats_score, get_ats_grade
                    ats_score_data = calculate_ats_score(resume_text)
                    ats_score_data['grade'] = get_ats_grade(ats_score_data['overall_score'])
        except Exception as e:
            print(f"Error calculating ATS score for existing resume: {e}")
            ats_score_data = None
    
    return templates.TemplateResponse("upload_resume.html", {
        "request": request,
        "message": None,
        "current_user": current_user,
        "is_supervisor": is_supervisor,
        "user": user,
        "ats_score": ats_score_data
    })

@app.post("/upload_resume", response_class=HTMLResponse)
async def upload_resume(request: Request, resume: UploadFile = Form(...), db: Session = Depends(get_db)):
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    
    # Save the uploaded file with username as filename
    file_ext = os.path.splitext(resume.filename)[1].lower()
    new_filename = f"{current_user}{file_ext}"
    file_location = os.path.join(RESUMES_DIR, new_filename)
    
    # Remove any previous resume files for this user (to avoid clutter and confusion)
    for ext in [".pdf", ".docx", ".doc"]:
        old_path = os.path.join(RESUMES_DIR, f"{current_user}{ext}")
        if os.path.exists(old_path) and old_path != file_location:
            os.remove(old_path)
    
    with open(file_location, "wb") as buffer:
        shutil.copyfileobj(resume.file, buffer)

    # Resume parsing
    resume_text = None
    error = None
    try:
        if file_ext == ".pdf":
            import pdfplumber
            with pdfplumber.open(file_location) as pdf:
                resume_text = " ".join(page.extract_text() or '' for page in pdf.pages)
        elif file_ext == ".docx":
            import docx
            doc = docx.Document(file_location)
            resume_text = "\n".join([para.text for para in doc.paragraphs])
        else:
            error = "Unsupported file type. Please upload a PDF or DOCX file."
    except Exception as e:
        error = f"Error parsing resume: {str(e)}"

    # Extract user info from resume_text (improved extraction)
    full_name = None
    education = None
    skills = None
    email = None
    phone = None
    try:
        from pyresparser import ResumeParser
        data = ResumeParser(file_location).get_extracted_data()
        # Always check the first non-empty, all-uppercase, single-word line as a possible name
        first_lines = resume_text.split('\n')[:20] if resume_text else []
        forced_name = None
        for line in first_lines:
            line = line.strip()
            if line and line.isupper() and 1 <= len(line.split()) <= 2 and line.replace(' ', '').isalpha():
                from .nlp_utils import SKILLS_TAXONOMY
                known_skills = set()
                for cat in SKILLS_TAXONOMY.values():
                    for skill in cat:
                        known_skills.add(skill.lower())
                if line.lower() not in known_skills:
                    forced_name = line.title()
                    break
        full_name = forced_name or data.get('name')
        
        # Filter out date patterns from education data
        education_data = data.get('education', [])
        filtered_education = []
        date_patterns = [
            r'\b(?:january|february|march|april|may|june|july|august|september|october|november|december)\s+\d{4}\s+(?:to|until|till|-)\s+(?:january|february|march|april|may|june|july|august|september|october|november|december)\s+\d{4}\b',
            r'\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\s+\d{4}\s+(?:to|until|till|-)\s+(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\s+\d{4}\b',
            r'\b\d{4}\s+(?:to|until|till|-)\s+\d{4}\b',
            r'\b(?:january|february|march|april|may|june|july|august|september|october|november|december)\s+\d{4}\b',
            r'\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\s+\d{4}\b',
            r'\b\d{1,2}/\d{1,2}/\d{4}\s+(?:to|until|till|-)\s+\d{1,2}/\d{1,2}/\d{4}\b',
            r'\b\d{1,2}-\d{1,2}-\d{4}\s+(?:to|until|till|-)\s+\d{1,2}-\d{1,2}-\d{4}\b'
        ]
        
        for edu_item in education_data:
            is_date_only = False
            for pattern in date_patterns:
                if re.search(pattern, edu_item, re.IGNORECASE):
                    is_date_only = True
                    break
            if not is_date_only:
                filtered_education.append(edu_item)
        
        education = ', '.join(filtered_education) if filtered_education else None
        skills = ', '.join(data.get('skills', [])) if data.get('skills') else None
        email = data.get('email')
        phone = data.get('mobile_number')
        print(f"DEBUG - PyResparser extracted: Name={full_name}, Education={education}, Skills={skills}, Email={email}, Phone={phone}")
    except Exception as e:
        print(f"DEBUG - Error in PyResparser: {e}")
        if resume_text:
            try:
                from .nlp_utils import extract_name, extract_education, extract_skills, extract_email, extract_phone, SKILLS_TAXONOMY
                # Always check the first non-empty, all-uppercase, single-word line as a possible name
                first_lines = resume_text.split('\n')[:20]
                forced_name = None
                known_skills = set()
                for cat in SKILLS_TAXONOMY.values():
                    for skill in cat:
                        known_skills.add(skill.lower())
                for line in first_lines:
                    line = line.strip()
                    if line and line.isupper() and 1 <= len(line.split()) <= 2 and line.replace(' ', '').isalpha():
                        if line.lower() not in known_skills:
                            forced_name = line.title()
                            break
                full_name = forced_name or extract_name(resume_text)
                education = extract_education(resume_text)
                skills_list = extract_skills(resume_text)
                skills = ", ".join([skill for category_skills in skills_list.values() for skill in category_skills]) if skills_list else None
                email = extract_email(resume_text)
                phone = extract_phone(resume_text)
                print(f"DEBUG - Fallback extracted: Name={full_name}, Education={education}, Skills={skills}, Email={email}, Phone={phone}")
            except Exception as e2:
                print(f"DEBUG - Error in fallback parsing: {e2}")
                lines = resume_text.split('\n')
                full_name = lines[0] if lines else None
                education = None
                skills = None
                email = None
                phone = None
    # Calculate ATS score
    ats_score_data = None
    if resume_text and not error:
        try:
            from .ats_utils import calculate_ats_score, get_ats_grade
            ats_score_data = calculate_ats_score(resume_text)
            ats_score_data['grade'] = get_ats_grade(ats_score_data['overall_score'])
        except Exception as e:
            print(f"Error calculating ATS score: {e}")
            ats_score_data = None
    
    # Update user profile
    user = db.query(models.User).filter(models.User.username == current_user).first()
    if user:
        user.resume_filename = new_filename
        user.full_name = full_name
        user.education = education
        user.skills = skills
        user.resume_email = email
        user.resume_phone = phone
        db.commit()
        db.refresh(user)
    
    message = "Resume uploaded and parsed successfully!"
    if error:
        message = error
    
    return templates.TemplateResponse("upload_resume.html", {
        "request": request, 
        "message": message, 
        "current_user": current_user, 
        "is_supervisor": is_supervisor,
        "user": user,
        "email": email,
        "phone": phone,
        "ats_score": ats_score_data
    })

@app.post("/delete_resume", response_class=HTMLResponse)
async def delete_resume(request: Request, db: Session = Depends(get_db)):
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    user = db.query(models.User).filter(models.User.username == current_user).first()
    if user and user.resume_filename:
        file_path = os.path.join(RESUMES_DIR, user.resume_filename)
        if os.path.exists(file_path):
            os.remove(file_path)
        user.resume_filename = None
        user.full_name = None
        user.education = None
        user.skills = None
        db.commit()
    message = "Resume deleted successfully."
    return templates.TemplateResponse("upload_resume.html", {
        "request": request, 
        "message": message, 
        "current_user": current_user, 
        "is_supervisor": is_supervisor,
        "user": user
    })

@app.get("/job/{job_id}")
async def view_job_details(
    request: Request,
    job_id: int,
    db: Session = Depends(get_db)
):
    """View full job details with complete description"""
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    
    job = db.query(models.Job).filter(models.Job.id == job_id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    
    # Check if user has already applied
    user = db.query(models.User).filter(models.User.username == current_user).first()
    existing_application = db.query(models.Application).filter(
        models.Application.job_id == job_id,
        models.Application.user_id == user.id
    ).first()
    
    has_applied = existing_application is not None
    
    # Add preview/full description for Read More toggle
    try:
        from .job_description_processor import create_structured_description
        structured = create_structured_description(job.description or "")
        setattr(job, 'description_preview', structured.get('preview', ''))
        setattr(job, 'has_more', structured.get('has_more', False))
    except Exception:
        setattr(job, 'description_preview', (job.description or '')[:400])
        setattr(job, 'has_more', bool(job.description and len(job.description) > 400))
    
    return templates.TemplateResponse("job_details.html", {
        "request": request,
        "job": job,
        "current_user": current_user,
        "is_supervisor": is_supervisor,
        "applied": has_applied,
        "user": user
    })

@app.get("/analyze_job_match/{job_id}")
async def analyze_job_match(
    request: Request,
    job_id: int,
    applied: bool = False,
    db: Session = Depends(get_db)
):
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    
    # Get the job posting
    job = db.query(models.Job).filter(models.Job.id == job_id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    
    # Look for user's resume in both PDF and DOCX formats
    resume_path = None
    for ext in ['.pdf', '.docx']:
        temp_path = os.path.join(RESUMES_DIR, f"{current_user}{ext}")
        if os.path.exists(temp_path):
            resume_path = temp_path
            break

    # Extra check: ensure the file actually exists
    if not resume_path or not os.path.isfile(resume_path):

        return templates.TemplateResponse(
            "error.html",
            {
                "request": request,
                "message": "Please upload your resume first (PDF or DOCX format)",
                "current_user": current_user,
                "is_supervisor": is_supervisor
            }
        )
    
    # Extract text from resume
    try:
        resume_text = ""
        if resume_path.endswith('.pdf'):
            import pdfplumber
            with pdfplumber.open(resume_path) as pdf:
                resume_text = " ".join(page.extract_text() or '' for page in pdf.pages)
        else:  # .docx
            import docx
            doc = docx.Document(resume_path)
            resume_text = "\n".join([para.text for para in doc.paragraphs])
        if not resume_text.strip():
            return templates.TemplateResponse(
                "error.html",
                {
                    "request": request,
                    "message": "Your resume could not be read. Please upload a valid PDF or DOCX file.",
                    "current_user": current_user,
                    "is_supervisor": is_supervisor
                }
            )
    except Exception as e:
        return templates.TemplateResponse(
            "error.html",
            {
                "request": request,
                "message": f"Error reading resume: {str(e)}",
                "current_user": current_user,
                "is_supervisor": is_supervisor
            }
        )
    
    # Extract skills from both job and resume
    job_text = (job.description or "") + " " + (job.skills or "")
    job_skills = extract_skills(job_text)
    resume_skills = extract_skills(resume_text)
    
    # Calculate match scores
    match_scores = calculate_skill_match(job_skills, resume_skills)
    # Get skill improvement suggestions
    suggestions = get_skill_suggestions(job_skills, resume_skills)

    overall_score = match_scores.get('overall', 0)
    
    # Calculate ATS score for this specific job
    ats_score_data = None
    try:
        from .ats_utils import calculate_ats_score, get_ats_grade
        job_description = (job.description or "") + " " + (job.skills or "")
        ats_score_data = calculate_ats_score(resume_text, job_description)
        ats_score_data['grade'] = get_ats_grade(ats_score_data['overall_score'])
    except Exception as e:
        print(f"Error calculating ATS score for job match: {e}")
        ats_score_data = None

    # NEW: Calculate matched and missing skills for each category
    matched_skills = {}
    missing_skills = {}
    for category in job_skills:
        job_set = set(job_skills.get(category, []))
        resume_set = set(resume_skills.get(category, []))
        matched_skills[category] = list(job_set & resume_set)
        missing_skills[category] = list(job_set - resume_set)

    # Format the results for display (improved ATS style)
    results_html = """
    <div style='padding: 1.5rem; background: #f4f8ff; border-radius: 1rem; box-shadow: 0 2px 8px #e0e7ef;'>
      <h2 style='font-size: 1.3rem; font-weight: bold; color: #1d3557; margin-bottom: 0.5rem;'>ATS Resume Score</h2>
      <div style='font-size: 1.1rem; margin-bottom: 0.5rem;'>
        <span style='color: #2563eb; font-weight: bold;'>Overall Match: {overall:.1f}%</span>
      </div>
      <div style='margin-bottom: 0.5rem;'><b>Category Scores:</b></div>
      <ul style='margin-bottom: 0.7rem; padding-left: 1.2rem;'>
    """
    results_html = results_html.format(overall=overall_score)
    for category, score in match_scores.items():
        if category != 'overall':
            results_html += f"<li><b>{category.replace('_', ' ').title()}:</b> <span style='color:#2563eb'>{score:.1f}%</span></li>"
    results_html += "</ul>"

    # Matched Skills
    matched_any = any(skills for skills in matched_skills.values())
    if matched_any:
        results_html += "<div style='margin-bottom: 0.3rem;'><b>Matched Skills:</b></div><ul style='margin-bottom: 0.7rem; padding-left: 1.2rem;'>"
        for category, skills in matched_skills.items():
            if skills:
                results_html += f"<li><b>{category.replace('_', ' ').title()}:</b> <span style='color:#059669'>{', '.join(skills)}</span></li>"
        results_html += "</ul>"

    # Missing Skills
    missing_any = any(skills for skills in missing_skills.values())
    if missing_any:
        results_html += "<div style='margin-bottom: 0.3rem;'><b>Missing Skills (to improve):</b></div><ul style='margin-bottom: 0.7rem; padding-left: 1.2rem;'>"
        for category, skills in missing_skills.items():
            if skills:
                results_html += f"<li><b>{category.replace('_', ' ').title()}:</b> <span style='color:#dc2626'>{', '.join(skills)}</span></li>"
        results_html += "</ul>"

    # Skill suggestions (optional, already covered by missing skills)
    if suggestions:
        results_html += "<div style='margin-bottom: 0.3rem;'><b>Suggested Skills to Learn:</b></div><ul style='margin-bottom: 0.7rem; padding-left: 1.2rem;'>"
        for category, missing in suggestions.items():
            if missing:
                results_html += f"<li><b>{category.replace('_', ' ').title()}:</b> <span style='color:#f59e42'>{', '.join(missing)}</span></li>"
        results_html += "</ul>"

    results_html += "</div>"
    
    user = db.query(models.User).filter(models.User.username == current_user).first()
    
    # Check if user has already applied for this job
    existing_application = db.query(models.Application).filter(
        models.Application.job_id == job_id,
        models.Application.user_id == user.id
    ).first()
    
    has_applied = existing_application is not None
    
    # Add preview/full description for Read More toggle
    try:
        from .job_description_processor import create_structured_description
        structured = create_structured_description(job.description or "")
        setattr(job, 'description_preview', structured.get('preview', ''))
        setattr(job, 'has_more', structured.get('has_more', False))
    except Exception:
        setattr(job, 'description_preview', (job.description or '')[:400])
        setattr(job, 'has_more', bool(job.description and len(job.description) > 400))
    
    return templates.TemplateResponse(
        "job_match.html",
        {
            "request": request,
            "job": job,
            "results": results_html,
            "current_user": current_user,
            "is_supervisor": is_supervisor,
            "applied": has_applied,
            "user": user,
            "ats_score": ats_score_data
        }
    )

@app.get("/my-applications", response_class=HTMLResponse)
async def my_applications(request: Request, db: Session = Depends(get_db)):
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    user = db.query(models.User).filter(models.User.username == current_user).first()
    applications = db.query(models.Application).filter(models.Application.user_id == user.id).all()
    jobs = db.query(models.Job).all()
    return templates.TemplateResponse("my_applications.html", {
        "request": request,
        "jobs": jobs,
        "applications": applications,
        "current_user": current_user,
        "is_supervisor": is_supervisor
    })

@app.get("/profile", response_class=HTMLResponse)
async def profile_page(request: Request, db: Session = Depends(get_db)):
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    user = db.query(models.User).filter(models.User.username == current_user).first()
    return templates.TemplateResponse("profile.html", {
        "request": request,
        "user": user,
        "current_user": current_user,
        "is_supervisor": is_supervisor,
        "message": None,
        "error": None
    })

@app.post("/profile", response_class=HTMLResponse)
async def update_profile(
    request: Request,
    full_name: str = Form(None),
    email: str = Form(None),
    resume_phone: str = Form(None),
    db: Session = Depends(get_db)
):
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    user = db.query(models.User).filter(models.User.username == current_user).first()
    if not user:
        return templates.TemplateResponse("profile.html", {
            "request": request,
            "user": user,
            "current_user": current_user,
            "is_supervisor": is_supervisor,
            "message": None,
            "error": "User not found."
        })
    # Check if email is being changed to one that already exists
    if email and email != user.email:
        existing = db.query(models.User).filter(models.User.email == email).first()
        if existing:
            return templates.TemplateResponse("profile.html", {
                "request": request,
                "user": user,
                "current_user": current_user,
                "is_supervisor": is_supervisor,
                "message": None,
                "error": "Email already in use."
            })
        user.email = email
    user.full_name = full_name
    user.resume_phone = resume_phone
    db.commit()
    db.refresh(user)
    return templates.TemplateResponse("profile.html", {
        "request": request,
        "user": user,
        "current_user": current_user,
        "is_supervisor": is_supervisor,
        "message": "Profile updated successfully.",
        "error": None
    })

@app.get("/my-reviews", response_class=HTMLResponse)
async def my_reviews(request: Request, db: Session = Depends(get_db)):
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    user = db.query(models.User).filter(models.User.username == current_user).first()
    reviews = db.query(models.Review).filter(models.Review.user_id == user.id).all()
    review_data = []
    for review in reviews:
        job = db.query(models.Job).filter(models.Job.id == review.job_id).first()
        application = db.query(models.Application).filter(
            models.Application.user_id == user.id,
            models.Application.job_id == review.job_id
        ).first()
        review_data.append({
            "id": review.id,  # Ensure id is included for delete
            "job": job,
            "rating": review.rating,
            "comment": review.comment,
            "created_at": review.created_at,
            "application_status": application.status if application else "N/A"
        })
    return templates.TemplateResponse("my_reviews.html", {
        "request": request,
        "reviews": review_data,
        "current_user": current_user,
        "is_supervisor": is_supervisor
    })

@app.post("/delete-review/{review_id}")
async def delete_review(review_id: int, db: Session = Depends(get_db)):
    if not current_user:
        raise HTTPException(status_code=401, detail="Not authenticated")
    
    review = db.query(models.Review).filter(models.Review.id == review_id).first()
    if not review:
        raise HTTPException(status_code=404, detail="Review not found")
    
    # Check if user owns this review
    user = db.query(models.User).filter(models.User.username == current_user).first()
    if review.user_id != user.id:
        raise HTTPException(status_code=403, detail="Not authorized")
    
    db.delete(review)
    db.commit()
    
    return {"message": "Review deleted successfully"}

@app.get("/api/notifications")
async def get_notifications(db: Session = Depends(get_db)):
    """Get user's notifications"""
    if not current_user:
        raise HTTPException(status_code=401, detail="Not authenticated")
    
    user = db.query(models.User).filter(models.User.username == current_user).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    notifications = db.query(models.Notification).filter(
        models.Notification.user_id == user.id
    ).order_by(models.Notification.created_at.desc()).limit(10).all()
    
    return {
        "notifications": [
            {
                "id": notif.id,
                "title": notif.title,
                "message": notif.message,
                "type": notif.notification_type,
                "is_read": notif.is_read,
                "created_at": notif.created_at.isoformat(),
                "related_job_id": notif.related_job_id
            }
            for notif in notifications
        ],
        "unread_count": len([n for n in notifications if not n.is_read])
    }

@app.post("/api/notifications/{notification_id}/read")
async def mark_notification_read(notification_id: int, db: Session = Depends(get_db)):
    """Mark a notification as read"""
    if not current_user:
        raise HTTPException(status_code=401, detail="Not authenticated")
    
    user = db.query(models.User).filter(models.User.username == current_user).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    notification = db.query(models.Notification).filter(
        models.Notification.id == notification_id,
        models.Notification.user_id == user.id
    ).first()
    
    if not notification:
        raise HTTPException(status_code=404, detail="Notification not found")
    
    notification.is_read = True
    db.commit()
    
    return {"message": "Notification marked as read"}

@app.post("/api/notifications/read-all")
async def mark_all_notifications_read(db: Session = Depends(get_db)):
    """Mark all user notifications as read"""
    if not current_user:
        raise HTTPException(status_code=401, detail="Not authenticated")
    
    user = db.query(models.User).filter(models.User.username == current_user).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    db.query(models.Notification).filter(
        models.Notification.user_id == user.id,
        models.Notification.is_read == False
    ).update({"is_read": True})
    
    db.commit()
    
    return {"message": "All notifications marked as read"}

# Resume Builder Routes
@app.get("/resume-builder", response_class=HTMLResponse)
async def resume_builder_page(request: Request):
    """Resume builder page"""
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    
    return templates.TemplateResponse("resume_builder.html", {
        "request": request,
        "current_user": current_user,
        "is_supervisor": is_supervisor
    })

@app.post("/create-ai-resume")
async def create_ai_resume(
    request: Request,
    resume_name: str = Form(...),
    target_role: str = Form(...),
    experience_level: str = Form(...),
    industry: str = Form(...),
    key_skills: str = Form(""),
    years_experience: str = Form(""),
    template: str = Form("professional"),
    user_prompt: str = Form(...),
    db: Session = Depends(get_db)
):
    """Create AI-generated resume"""
    if not current_user:
        raise HTTPException(status_code=401, detail="Not authenticated")
    
    try:
        # Get user
        user = db.query(models.User).filter(models.User.username == current_user).first()
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        
        # Prepare user data
        user_data = {
            "target_role": target_role,
            "experience_level": experience_level,
            "industry": industry,
            "key_skills": key_skills,
            "years_experience": years_experience
        }
        
        # Create resume builder instance
        resume_builder = ResumeBuilder()
        
        # Generate resume
        result = resume_builder.create_resume(
            user_id=user.id,
            resume_name=resume_name,
            user_prompt=user_prompt,
            user_data=user_data,
            template_name=template,
            db=db
        )
        
        if result["success"]:
            return result
        else:
            return {"success": False, "message": result["message"]}
            
    except Exception as e:
        print(f"Error creating AI resume: {e}")
        return {"success": False, "message": f"Failed to create resume: {str(e)}"}

@app.get("/my-resumes", response_class=HTMLResponse)
async def my_resumes_page(request: Request, db: Session = Depends(get_db)):
    """My resumes page"""
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    
    try:
        # Get user
        user = db.query(models.User).filter(models.User.username == current_user).first()
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        
        # Get user's resumes
        resume_builder = ResumeBuilder()
        resumes = resume_builder.get_user_resumes(user.id, db)
        
        return templates.TemplateResponse("my_resumes.html", {
            "request": request,
            "resumes": resumes,
            "current_user": current_user,
            "is_supervisor": is_supervisor
        })
        
    except Exception as e:
        print(f"Error loading resumes: {e}")
        return templates.TemplateResponse("my_resumes.html", {
            "request": request,
            "resumes": [],
            "current_user": current_user,
            "is_supervisor": is_supervisor
        })

@app.get("/resume-preview/{resume_id}", response_class=HTMLResponse)
async def resume_preview_page(request: Request, resume_id: int, db: Session = Depends(get_db)):
    """Resume preview page"""
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    
    try:
        # Get user
        user = db.query(models.User).filter(models.User.username == current_user).first()
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        
        # Get resume
        resume_builder = ResumeBuilder()
        resume = resume_builder.get_resume_by_id(resume_id, user.id, db)
        
        if not resume:
            raise HTTPException(status_code=404, detail="Resume not found")
        
        return templates.TemplateResponse("resume_preview.html", {
            "request": request,
            "resume": resume,
            "current_user": current_user,
            "is_supervisor": is_supervisor
        })
        
    except Exception as e:
        print(f"Error loading resume preview: {e}")
        raise HTTPException(status_code=500, detail="Failed to load resume")

@app.get("/download-resume/{resume_id}")
async def download_resume(resume_id: int, db: Session = Depends(get_db)):
    """Download resume PDF"""
    if not current_user:
        raise HTTPException(status_code=401, detail="Not authenticated")
    
    try:
        # Get user
        user = db.query(models.User).filter(models.User.username == current_user).first()
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        
        # Get resume
        resume = db.query(GeneratedResume).filter(
            GeneratedResume.id == resume_id,
            GeneratedResume.user_id == user.id,
            GeneratedResume.is_active == True
        ).first()
        
        if not resume:
            raise HTTPException(status_code=404, detail="Resume not found")
        
        # Generate PDF on-demand
        resume_content = json.loads(resume.generated_content) if resume.generated_content else {}
        pdf_generator = ResumePDFGenerator()
        
        # Create temp directory if it doesn't exist
        os.makedirs("temp", exist_ok=True)
        
        # Create temporary PDF
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        pdf_filename = f"resume_{resume_id}_{timestamp}.pdf"
        pdf_path = pdf_generator.generate_pdf(
            resume_content, 
            resume.template_used, 
            f"temp/{pdf_filename}"
        )
        
        # Read and return the PDF
        with open(pdf_path, "rb") as f:
            content = f.read()
        
        # Clean up temporary file
        try:
            os.remove(pdf_path)
        except:
            pass
        
        return Response(
            content=content,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={resume.resume_name.replace(' ', '_')}.pdf"}
        )
        
    except Exception as e:
        print(f"Error downloading resume: {e}")
        raise HTTPException(status_code=500, detail="Failed to download resume")

@app.post("/delete-resume/{resume_id}")
async def delete_resume_route(resume_id: int, db: Session = Depends(get_db)):
    """Delete resume"""
    if not current_user:
        raise HTTPException(status_code=401, detail="Not authenticated")
    
    try:
        # Get user
        user = db.query(models.User).filter(models.User.username == current_user).first()
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        
        # Delete resume
        resume_builder = ResumeBuilder()
        success = resume_builder.delete_resume(resume_id, user.id, db)
        
        if success:
            return {"success": True, "message": "Resume deleted successfully"}
        else:
            raise HTTPException(status_code=404, detail="Resume not found")
            
    except Exception as e:
        print(f"Error deleting resume: {e}")
        raise HTTPException(status_code=500, detail="Failed to delete resume")

@app.get("/edit-resume/{resume_id}", response_class=HTMLResponse)
async def edit_resume_page(request: Request, resume_id: int, db: Session = Depends(get_db)):
    """Page to edit resume personal information"""
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    
    # Get the user object
    user = db.query(models.User).filter(models.User.username == current_user).first()
    if not user:
        return RedirectResponse(url="/login", status_code=303)
    
    # Get the resume
    resume = db.query(GeneratedResume).filter(
        GeneratedResume.id == resume_id,
        GeneratedResume.user_id == user.id,
        GeneratedResume.is_active == True
    ).first()
    
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    # Parse the resume content
    resume_content = json.loads(resume.generated_content) if resume.generated_content else {}
    personal_info = resume_content.get('personal_info', {})
    
    return templates.TemplateResponse("edit_resume.html", {
        "request": request,
        "current_user": current_user,
        "is_supervisor": is_supervisor,
        "resume": resume,
        "personal_info": personal_info
    })

@app.post("/update-resume/{resume_id}")
async def update_resume_route(
    request: Request,
    resume_id: int,
    name: str = Form(...),
    email: str = Form(...),
    phone: str = Form(...),
    location: str = Form(...),
    linkedin: str = Form(""),
    db: Session = Depends(get_db)
):
    """Update resume personal information"""
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)
    
    # Get the user object
    user = db.query(models.User).filter(models.User.username == current_user).first()
    if not user:
        return RedirectResponse(url="/login", status_code=303)
    
    # Get the resume
    resume = db.query(GeneratedResume).filter(
        GeneratedResume.id == resume_id,
        GeneratedResume.user_id == user.id,
        GeneratedResume.is_active == True
    ).first()
    
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    # Parse the current resume content
    resume_content = json.loads(resume.generated_content) if resume.generated_content else {}
    
    # Update personal information
    if 'personal_info' not in resume_content:
        resume_content['personal_info'] = {}
    
    resume_content['personal_info'].update({
        'name': name,
        'email': email,
        'phone': phone,
        'location': location,
        'linkedin': linkedin
    })
    
    # Save updated content
    resume.generated_content = json.dumps(resume_content)
    db.commit()
    
    return RedirectResponse(url=f"/resume-preview/{resume_id}", status_code=303)

@app.post("/regenerate-resume/{resume_id}")
async def regenerate_resume_route(
    request: Request,
    resume_id: int,
    new_prompt: str = Form(...),
    template: str = Form("professional"),
    db: Session = Depends(get_db)
):
    """Regenerate resume with new prompt"""
    if not current_user:
        raise HTTPException(status_code=401, detail="Not authenticated")
    
    try:
        # Get user
        user = db.query(models.User).filter(models.User.username == current_user).first()
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        
        # Prepare user data (you might want to get this from the original resume)
        user_data = {
            "target_role": "General Professional",
            "experience_level": "Mid-level",
            "industry": "Technology",
            "key_skills": "",
            "years_experience": "3-5"
        }
        
        # Regenerate resume
        resume_builder = ResumeBuilder()
        result = resume_builder.regenerate_resume(
            resume_id=resume_id,
            user_id=user.id,
            new_prompt=new_prompt,
            user_data=user_data,
            template_name=template,
            db=db
        )
        
        if result["success"]:
            return result
        else:
            raise HTTPException(status_code=500, detail=result["message"])
            
    except Exception as e:
        print(f"Error regenerating resume: {e}")
        raise HTTPException(status_code=500, detail="Failed to regenerate resume")